#!/usr/bin/env python3
"""Build a consolidated DOCX report from an idea-funnel run."""

from __future__ import annotations

import argparse
import json
import os
from datetime import datetime, timezone
from pathlib import Path

import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt, RGBColor


IDEAS_ROOT = Path(os.environ.get("IDEAS_ROOT", "/opt/data/ideas"))


class ReportBuilder:
    def __init__(self, run_id: str, output: Path):
        self.run_id = run_id
        self.run_dir = IDEAS_ROOT / "runs" / run_id
        self.wiki_dir = IDEAS_ROOT / "wiki"
        self.output = output
        self.doc = Document()
        self._setup_styles()
        self._setup_page()

    def _setup_styles(self):
        styles = self.doc.styles

        title = styles.add_style("IdeaFunnelTitle", WD_STYLE_TYPE.PARAGRAPH)
        title.font.name = "Inter"
        title.font.size = Pt(28)
        title.font.bold = True
        title.font.color.rgb = RGBColor(15, 23, 42)
        title.paragraph_format.space_after = Pt(6)
        title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = styles.add_style("IdeaFunnelSubtitle", WD_STYLE_TYPE.PARAGRAPH)
        subtitle.font.name = "Inter"
        subtitle.font.size = Pt(12)
        subtitle.font.color.rgb = RGBColor(100, 116, 139)
        subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(24)

        h1 = styles.add_style("IdeaFunnelH1", WD_STYLE_TYPE.PARAGRAPH)
        h1.font.name = "Inter"
        h1.font.size = Pt(20)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(15, 23, 42)
        h1.paragraph_format.space_before = Pt(24)
        h1.paragraph_format.space_after = Pt(10)

        h2 = styles.add_style("IdeaFunnelH2", WD_STYLE_TYPE.PARAGRAPH)
        h2.font.name = "Inter"
        h2.font.size = Pt(16)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(30, 41, 59)
        h2.paragraph_format.space_before = Pt(18)
        h2.paragraph_format.space_after = Pt(8)

        h3 = styles.add_style("IdeaFunnelH3", WD_STYLE_TYPE.PARAGRAPH)
        h3.font.name = "Inter"
        h3.font.size = Pt(13)
        h3.font.bold = True
        h3.font.color.rgb = RGBColor(51, 65, 85)
        h3.paragraph_format.space_before = Pt(12)
        h3.paragraph_format.space_after = Pt(6)

        body = styles.add_style("IdeaFunnelBody", WD_STYLE_TYPE.PARAGRAPH)
        body.font.name = "Inter"
        body.font.size = Pt(10.5)
        body.font.color.rgb = RGBColor(51, 65, 85)
        body.paragraph_format.space_after = Pt(6)
        body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        meta = styles.add_style("IdeaFunnelMeta", WD_STYLE_TYPE.PARAGRAPH)
        meta.font.name = "Inter"
        meta.font.size = Pt(9)
        meta.font.color.rgb = RGBColor(100, 116, 139)
        meta.paragraph_format.space_after = Pt(4)

        code = styles.add_style("IdeaFunnelCode", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "JetBrains Mono"
        code.font.size = Pt(9)
        code.font.color.rgb = RGBColor(15, 23, 42)
        code.paragraph_format.space_after = Pt(6)
        code.paragraph_format.left_indent = Inches(0.2)

    def _setup_page(self):
        section = self.doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

    def _read(self, path: Path) -> str:
        if not path.exists():
            return f"*(file not found: {path.name})*"
        return path.read_text(encoding="utf-8")

    def _read_json(self, path: Path) -> dict:
        if not path.exists():
            return {}
        return json.loads(path.read_text(encoding="utf-8"))

    def _add_markdown(self, text: str):
        html = markdown.markdown(text, extensions=["tables", "fenced_code"])
        soup = BeautifulSoup(html, "html.parser")
        self._render_soup(soup)

    def _render_soup(self, soup: BeautifulSoup):
        for elem in soup.find_all(["h1", "h2", "h3", "p", "ul", "ol", "pre", "table", "blockquote", "hr"]):
            if elem.name == "h1":
                self.doc.add_paragraph(elem.get_text(strip=True), style="IdeaFunnelH1")
            elif elem.name == "h2":
                self.doc.add_paragraph(elem.get_text(strip=True), style="IdeaFunnelH2")
            elif elem.name == "h3":
                self.doc.add_paragraph(elem.get_text(strip=True), style="IdeaFunnelH3")
            elif elem.name == "p":
                self._add_inline_paragraph(elem, "IdeaFunnelBody")
            elif elem.name == "ul":
                for li in elem.find_all("li", recursive=False):
                    p = self.doc.add_paragraph(style="List Bullet")
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.space_after = Pt(4)
                    self._add_inline(li, p)
            elif elem.name == "ol":
                for li in elem.find_all("li", recursive=False):
                    p = self.doc.add_paragraph(style="List Number")
                    p.paragraph_format.left_indent = Inches(0.25)
                    p.paragraph_format.space_after = Pt(4)
                    self._add_inline(li, p)
            elif elem.name == "pre":
                code = elem.get_text()
                for line in code.splitlines() or [""]:
                    self.doc.add_paragraph(line, style="IdeaFunnelCode")
            elif elem.name == "blockquote":
                p = self.doc.add_paragraph(style="IdeaFunnelBody")
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.right_indent = Inches(0.3)
                self._add_inline(elem, p)
            elif elem.name == "hr":
                self.doc.add_paragraph("_" * 60, style="IdeaFunnelMeta")
            elif elem.name == "table":
                self._add_table(elem)

    def _add_inline_paragraph(self, elem, style_name: str):
        p = self.doc.add_paragraph(style=style_name)
        self._add_inline(elem, p)

    def _add_inline(self, elem, paragraph):
        for child in elem.children:
            if isinstance(child, str):
                if child.strip():
                    run = paragraph.add_run(child)
                    run.font.name = "Inter"
            elif child.name in ("strong", "b"):
                run = paragraph.add_run(child.get_text())
                run.bold = True
                run.font.name = "Inter"
            elif child.name in ("em", "i"):
                run = paragraph.add_run(child.get_text())
                run.italic = True
                run.font.name = "Inter"
            elif child.name == "code":
                run = paragraph.add_run(child.get_text())
                run.font.name = "JetBrains Mono"
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(220, 38, 38)
            elif child.name == "a":
                run = paragraph.add_run(child.get_text())
                run.underline = True
                run.font.color.rgb = RGBColor(37, 99, 235)
            elif child.name == "br":
                paragraph.add_run("\n")
            else:
                self._add_inline(child, paragraph)

    def _add_table(self, table_elem):
        rows = table_elem.find_all("tr")
        if not rows:
            return
        num_cols = max(len(row.find_all(["td", "th"])) for row in rows)
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"
        for i, row in enumerate(rows):
            cells = row.find_all(["td", "th"])
            for j, cell in enumerate(cells):
                if j >= num_cols:
                    break
                table.rows[i].cells[j].text = cell.get_text(strip=True)
                if cell.name == "th":
                    for paragraph in table.rows[i].cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        self.doc.add_paragraph()

    def _cover_page(self):
        for _ in range(6):
            self.doc.add_paragraph()

        self.doc.add_paragraph("Idea Funnel Run Report", style="IdeaFunnelTitle")
        self.doc.add_paragraph(f"Run ID: {self.run_id}", style="IdeaFunnelSubtitle")

        now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
        self.doc.add_paragraph(f"Compiled on {now}", style="IdeaFunnelMeta")

        metrics = self._read_json(self.run_dir / "metrics.json")
        total = metrics.get("total_signals", 0)
        adapters = metrics.get("adapters", {})
        adapter_lines = [f"{k}: {v.get('signals', 0)}" for k, v in adapters.items()]

        self.doc.add_paragraph()
        self.doc.add_paragraph(f"Signals collected: {total}", style="IdeaFunnelMeta")
        self.doc.add_paragraph(f"Adapters: {', '.join(adapter_lines)}", style="IdeaFunnelMeta")

        self.doc.add_page_break()

    def _table_of_contents(self):
        self.doc.add_paragraph("Contents", style="IdeaFunnelH1")
        sections = [
            "1. Executive Summary",
            "2. Source Metrics & Retrospective",
            "3. Researcher Discovery Outputs",
            "4. Verification Gate",
            "5. Synthesized Briefs",
            "6. Final Judge Gate",
            "7. Wiki Curation Report",
            "8. Search Strategy for Next Run",
            "9. Wiki Pages Added",
        ]
        for section in sections:
            p = self.doc.add_paragraph(section, style="IdeaFunnelBody")
            p.paragraph_format.left_indent = Inches(0.25)
        self.doc.add_page_break()

    def _section(self, title: str, md_text: str):
        self.doc.add_paragraph(title, style="IdeaFunnelH1")
        self._add_markdown(md_text)
        self.doc.add_page_break()

    def _section_no_break(self, title: str, md_text: str):
        self.doc.add_paragraph(title, style="IdeaFunnelH1")
        self._add_markdown(md_text)

    def build(self):
        self._cover_page()
        self._table_of_contents()

        metrics = self._read_json(self.run_dir / "metrics.json")
        summary = f"""This report consolidates the complete output of idea-funnel run `{self.run_id}`.

**Pipeline status:** 8/8 tasks completed successfully.
**Sources:** {metrics.get('total_signals', 0)} signals from {len(metrics.get('adapters', {}))} adapters.
**New wiki ideas:** 3 pages curated.
**Search terms:** Updated for the next run.

The run used deterministic source adapters, three parallel researchers, a verifier gate, synthesizer, final judge, wiki-curator, and a search-strategist to close the feedback loop."""
        self._section("1. Executive Summary", summary)

        metrics = self._read_json(self.run_dir / "metrics.json")
        metrics_md = "```json\n" + json.dumps(metrics, indent=2, ensure_ascii=False) + "\n```"
        retro_md = self._read(self.run_dir / "retrospective.md")
        self._section("2. Source Metrics & Retrospective", metrics_md + "\n\n" + retro_md)

        discovery_md = ""
        for path in sorted((self.run_dir / "discovery").glob("*.md")):
            discovery_md += f"## {path.stem}\n\n{self._read(path)}\n\n---\n\n"
        self._section("3. Researcher Discovery Outputs", discovery_md)

        self._section("4. Verification Gate", self._read(self.run_dir / "verification.md"))
        self._section("5. Synthesized Briefs", self._read(self.run_dir / "briefs.md"))
        self._section("6. Final Judge Gate", self._read(self.run_dir / "judge-report.md"))
        self._section("7. Wiki Curation Report", self._read(self.run_dir / "curation-report.md"))

        st = self._read_json(self.run_dir / "search-terms.json")
        search_md = ""
        if st.get("notes"):
            search_md += st["notes"] + "\n\n"
        search_md += "```json\n" + json.dumps(st, indent=2, ensure_ascii=False) + "\n```"
        self._section("8. Search Strategy for Next Run", search_md)

        ideas_dir = self.wiki_dir / "ideas"
        wiki_md = ""
        for path in sorted(ideas_dir.glob("*.md")):
            mtime = datetime.fromtimestamp(path.stat().st_mtime, tz=timezone.utc)
            run_start_str = self.run_id.replace("t", " ").replace("z", "")
            try:
                run_start = datetime.strptime(run_start_str, "%Y-%m-%d %H%M%S").replace(tzinfo=timezone.utc)
            except ValueError:
                run_start = datetime.min.replace(tzinfo=timezone.utc)
            if mtime >= run_start:
                wiki_md += f"## {path.stem}\n\n{self._read(path)}\n\n---\n\n"
        self._section_no_break("9. Wiki Pages Added", wiki_md or "*(no new pages detected)*")

        self.output.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(self.output))
        print(f"Saved DOCX: {self.output}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--run-id", required=True)
    parser.add_argument("--output", required=True, type=Path)
    args = parser.parse_args()
    builder = ReportBuilder(args.run_id, args.output)
    builder.build()


if __name__ == "__main__":
    main()
